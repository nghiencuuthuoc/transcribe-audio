import os
from docx import Document

# === CONFIG ===
INPUT_FOLDER = "../mp3-test"

def convert_txt_to_docx(txt_path):
    with open(txt_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    doc.add_heading(f"Transcription: {os.path.basename(txt_path)}", level=1)

    for line in lines:
        if line.strip():
            doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph("")  # giữ dòng trống

    docx_path = os.path.splitext(txt_path)[0] + ".docx"
    doc.save(docx_path)
    print(f"✅ Saved: {docx_path}")

# === Loop all .txt files in all subfolders
for dirpath, _, filenames in os.walk(INPUT_FOLDER):
    for file in filenames:
        if file.endswith(".txt"):
            txt_file = os.path.join(dirpath, file)
            convert_txt_to_docx(txt_file)
